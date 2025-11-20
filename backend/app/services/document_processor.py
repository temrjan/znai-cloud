"""Document processing and indexing service using Llama Index."""
from typing import List
from pathlib import Path
import os

from llama_index.core import Document, VectorStoreIndex, Settings, StorageContext
from llama_index.core.vector_stores import MetadataFilters, ExactMatchFilter
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms.openai import OpenAI
from llama_index.vector_stores.qdrant import QdrantVectorStore
from llama_index.core.node_parser import SentenceSplitter
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams

from backend.app.config import settings


class DocumentProcessor:
    """Process and index documents using Llama Index and Qdrant."""

    def __init__(self):
        # Setup OpenAI embedding model
        self.embed_model = OpenAIEmbedding(
            model="text-embedding-3-large",
            api_key=settings.openai_api_key,
            dimensions=3072
        )

        # Setup OpenAI LLM
        self.llm = OpenAI(
            model="gpt-4o",
            api_key=settings.openai_api_key,
        )

        # Set global Llama Index settings
        Settings.embed_model = self.embed_model
        Settings.llm = self.llm
        Settings.chunk_size = 512  # Like in biotact_core
        Settings.chunk_overlap = 50  # Like in biotact_core

        # Setup Qdrant client
        self.qdrant_client = QdrantClient(
            host=settings.qdrant_host,
            port=settings.qdrant_port,
        )

        self.collection_name = "ai_avangard_documents"

        # Setup sentence splitter for chunking
        self.text_splitter = SentenceSplitter(
            chunk_size=512,
            chunk_overlap=50,
        )

        self._ensure_collection()

    def _ensure_collection(self):
        """Create collection if it doesn't exist."""
        collections = self.qdrant_client.get_collections().collections
        if not any(c.name == self.collection_name for c in collections):
            self.qdrant_client.create_collection(
                collection_name=self.collection_name,
                vectors_config=VectorParams(
                    size=3072,  # text-embedding-3-large dimension
                    distance=Distance.COSINE,
                ),
            )

    def extract_text(self, file_path: Path, mime_type: str) -> str:
        """Extract text from various file formats."""
        text = ""

        if mime_type == "text/plain" or file_path.suffix in [".txt", ".md"]:
            text = file_path.read_text(encoding="utf-8", errors="ignore")

        elif mime_type == "application/pdf" or file_path.suffix == ".pdf":
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = "\n".join(page.extract_text() for page in reader.pages)
            except Exception as e:
                raise ValueError(f"Failed to extract PDF: {str(e)}")

        elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"] or file_path.suffix in [".doc", ".docx"]:
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            except Exception as e:
                raise ValueError(f"Failed to extract DOCX: {str(e)}")

        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

        return text.strip()

    def detect_content_type(self, text: str, filename: str) -> str:
        """Detect content type using simple heuristics."""
        text_lower = text.lower()
        filename_lower = filename.lower()

        # Legal keywords
        legal_keywords = ['договор', 'соглашение', 'закон', 'статья', 'пункт', 'постановление',
                         'contract', 'agreement', 'law', 'article', 'clause', 'hereby']
        if any(kw in text_lower for kw in legal_keywords[:6]) or \
           any(kw in filename_lower for kw in ['contract', 'legal', 'agreement', 'договор']):
            return 'legal'

        # Technical keywords
        tech_keywords = ['api', 'function', 'class', 'алгоритм', 'код', 'программа',
                        'database', 'algorithm', 'implementation', 'method']
        if any(kw in text_lower for kw in tech_keywords[:6]) or \
           any(kw in filename_lower for kw in ['tech', 'api', 'code', 'dev']):
            return 'technical'

        # Cooking keywords
        cooking_keywords = ['рецепт', 'ингредиент', 'приготовление', 'духовка', 'минут',
                           'recipe', 'ingredient', 'cooking', 'bake', 'minutes', 'tablespoon']
        if any(kw in text_lower for kw in cooking_keywords[:5]) or \
           any(kw in filename_lower for kw in ['recipe', 'cooking', 'рецепт']):
            return 'cooking'

        return 'general'

    def index_document(
        self,
        document_id: str,
        user_id: int,
        filename: str,
        file_path: Path,
        mime_type: str,
    ) -> int:
        """
        Process and index document using Llama Index.

        Returns number of chunks indexed.
        """
        # Extract text
        text = self.extract_text(file_path, mime_type)

        if not text:
            raise ValueError("No text extracted from document")

        # Detect content type
        content_type = self.detect_content_type(text, filename)

        # Create Llama Index Document with metadata
        doc = Document(
            text=text,
            metadata={
                "document_id": str(document_id),
                "user_id": user_id,
                "filename": filename,
                "content_type": content_type,
            }
        )

        # Create vector store
        vector_store = QdrantVectorStore(
            client=self.qdrant_client,
            collection_name=self.collection_name,
        )

        # Create storage context
        storage_context = StorageContext.from_defaults(
            vector_store=vector_store
        )

        # Create index with automatic chunking using Llama Index
        index = VectorStoreIndex.from_documents(
            [doc],
            storage_context=storage_context,
            transformations=[self.text_splitter],
            show_progress=False
        )

        # Count chunks by getting nodes
        nodes = self.text_splitter.get_nodes_from_documents([doc])

        return len(nodes)

    def delete_document(self, document_id: str):
        """Delete all chunks for a document from Qdrant."""
        # Note: Llama Index doesn't provide direct delete by metadata
        # We need to use Qdrant client directly
        try:
            self.qdrant_client.delete(
                collection_name=self.collection_name,
                points_selector={
                    "filter": {
                        "must": [
                            {"key": "document_id", "match": {"value": str(document_id)}}
                        ]
                    }
                },
            )
        except Exception as e:
            print(f"Warning: Failed to delete from Qdrant: {e}")

    def search(self, user_id: int, query: str, limit: int = 5, score_threshold: float = 0.35) -> List[dict]:
        """
        Search for relevant chunks using Llama Index.

        Args:
            user_id: User ID for filtering
            query: Search query
            limit: Maximum number of results
            score_threshold: Minimum similarity score (0.0-1.0), default 0.35

        Returns list of matching chunks with metadata.
        """
        # Create vector store
        vector_store = QdrantVectorStore(
            client=self.qdrant_client,
            collection_name=self.collection_name,
        )

        # Create index from existing vector store
        index = VectorStoreIndex.from_vector_store(
            vector_store=vector_store,
        )

        # Create query engine with user_id filter
        filters = MetadataFilters(
            filters=[ExactMatchFilter(key="user_id", value=user_id)]
        )

        query_engine = index.as_query_engine(
            similarity_top_k=limit,
            filters=filters
        )

        # Execute query
        response = query_engine.query(query)

        # Extract results and filter by score threshold
        results = []
        if hasattr(response, 'source_nodes'):
            for node in response.source_nodes:
                score = node.score if hasattr(node, 'score') else 0.0

                # Filter out results below threshold
                if score >= score_threshold:
                    results.append({
                        "text": node.text,
                        "filename": node.metadata.get("filename", "Unknown"),
                        "score": score,
                    })

        return results


# Singleton instance
document_processor = DocumentProcessor()
